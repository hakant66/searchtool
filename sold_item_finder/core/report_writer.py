from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches

from sold_item_finder.core.text_search import TextSearchResult
from sold_item_finder.core.email.models import EmailMessage


class ReportWriter:
    def __init__(self, base_dir: Path | None = None) -> None:
        self.base_dir = base_dir or (Path.home() / "Documents" / "SoldItemFinder" / "Reports")

    def write_email_report(
        self,
        email_message: EmailMessage,
        query: str,
        results: list[TextSearchResult],
        index_version: str = "v1",
    ) -> Path:
        day_dir = self.base_dir / datetime.now().strftime("%Y-%m-%d")
        day_dir.mkdir(parents=True, exist_ok=True)
        safe_subject = sanitize_filename(email_message.subject or "no_subject")
        safe_msg_id = sanitize_filename(email_message.message_id or "no_msgid")
        out_path = day_dir / f"{safe_subject}_{safe_msg_id}.docx"

        doc = Document()
        doc.add_heading("Sold Item Match Report", level=1)
        doc.add_paragraph(f"Date: {email_message.date}")
        doc.add_paragraph(f"From: {email_message.from_address}")
        doc.add_paragraph(f"To: {email_message.to_address}")
        doc.add_paragraph(f"Subject: {email_message.subject}")
        doc.add_paragraph(f"Search query: {query}")

        doc.add_heading("Top Results", level=2)
        for idx, result in enumerate(results[:10], start=1):
            doc.add_heading(f"{idx}. {result.title}", level=3)
            doc.add_paragraph(f"Platform: {result.platform}")
            doc.add_paragraph(f"SKU: {result.sku}")
            doc.add_paragraph(f"Listing ID: {result.listing_id}")
            doc.add_paragraph(f"Notes: {result.notes}")
            doc.add_paragraph(f"Folder: {result.path}")
            doc.add_paragraph(f"Relevance score: {result.score:.4f}")
            for image in result.image_paths[:2]:
                image_path = Path(image)
                if image_path.exists():
                    try:
                        doc.add_picture(str(image_path), width=Inches(2.2))
                    except Exception:
                        pass

        doc.add_paragraph(
            f"Generated: {datetime.now().isoformat(timespec='seconds')} | Index version: {index_version}"
        )
        doc.save(out_path)
        return out_path


def sanitize_filename(value: str) -> str:
    return re.sub(r"[^a-zA-Z0-9._-]+", "_", value).strip("._")[:90] or "report"
